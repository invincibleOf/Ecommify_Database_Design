# -*- coding: utf-8 -*-
# Generador del documento técnico ensamblado (Entregable 2 · Ecommify · E01).
# Alineado con "Implementación técnica completa en PostgreSQL y MongoDB V2.pdf".
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = "/Users/carlosbraydi/Documents/Maestria Arquitectura de Software/Optimización de bases de datos/Ecommify_Database_Design"
EVID = os.path.join(REPO, "evidencias")
EXPL = os.path.join(EVID, "explain_postgresql")
EJ3  = os.path.join(REPO, "..", "Ejercicio3")
OUT  = os.path.join(REPO, "docs", "Documento_Tecnico_Ecommify_E01.docx")

VIDEO_URL = "https://unisabanaedu-my.sharepoint.com/:b:/g/personal/camilomorro_unisabana_edu_co/IQBuKV8Mi-fMQ5uJjT3bZboMATD5nDk7ngv-_4QOJ6WY4-U?e=7bvxnd"

AZUL = RGBColor(0x1f,0x36,0x6a)
GRIS = RGBColor(0x55,0x55,0x55)

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

def H(text, lvl=1):
    h = doc.add_heading(text, level=lvl)
    for r in h.runs: r.font.color.rgb = AZUL
    return h

def P(text="", bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(9)
        run.font.color.rgb = RGBColor(0xff,0xff,0xff)
        sh = OxmlElement("w:shd"); sh.set(qn("w:fill"),"1f366a")
        c._tc.get_or_add_tcPr().append(sh)
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""; run=cells[i].paragraphs[0].add_run(str(val)); run.font.size=Pt(9)
    doc.add_paragraph()
    return t

def code(text):
    """Bloque de código monoespaciado con fondo gris claro."""
    t = doc.add_table(rows=1, cols=1); t.style="Table Grid"
    cell = t.rows[0].cells[0]
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"),"f4f4f4")
    cell._tc.get_or_add_tcPr().append(sh)
    cell.text=""
    for k,line in enumerate(text.strip("\n").split("\n")):
        p = cell.paragraphs[0] if k==0 else cell.add_paragraph()
        r = p.add_run(line if line else " "); r.font.name="Consolas"; r.font.size=Pt(8.5)
    doc.add_paragraph()

def figura(path, caption, width=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = P(caption, italic=True, size=9, color=GRIS); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    else:
        P("[Figura no encontrada: "+os.path.basename(path)+"]", italic=True, color=GRIS, size=9)

# ===================== PORTADA =====================
title = doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=title.add_run("Implementación técnica completa en PostgreSQL y MongoDB"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=AZUL
sub = doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run("Entregable 2 — Proyecto integrador Ecommify · Arquitectura políglota"); r.font.size=Pt(13); r.italic=True
m = doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
m.add_run("Dataset: Brazilian E-commerce (Olist) · Asignatura EFMAS-DOBD20263 · Junio 2026").font.size=Pt(11)
doc.add_paragraph()
P("Equipo de trabajo: E01", bold=True)
table(["Integrante","Código"],[
 ["Camilo José Mora Rodríguez","0000394950"],
 ["Carlos Alberto Zambrano Braydi","0000395349"],
 ["Alejandro Miguel López Castañeda","0000385228"],
 ["Carlos Arturo Salazar Castañeda","0000393238"]])
P("Entorno: PostgreSQL 17 sobre Supabase · MongoDB Atlas. Mediciones reales con "
  "EXPLAIN (ANALYZE, BUFFERS) y .explain(\"executionStats\").", italic=True, size=9, color=GRIS)
doc.add_page_break()

# ===================== a =====================
H("a. Resumen ejecutivo")
P("Ecommify implementa una arquitectura políglota que combina PostgreSQL (núcleo transaccional) "
  "y MongoDB (módulo analítico y de catálogo), optimizada para reportes de ventas, análisis de "
  "clientes y operación comercial sobre datos reales del dataset Olist (100.000 órdenes).")
P("Principales optimizaciones aplicadas:", bold=True)
for b in ["PostgreSQL: índices especializados (B-tree, parcial, BRIN, covering con INCLUDE), particionamiento por RANGE y reescritura de consultas críticas (CTEs).",
          "MongoDB: patrones de documento (Attribute/Embedding, Computed, Referenced), índices con regla ESR, índice parcial e índice de texto, y aggregation pipeline de 7 stages."]:
    bullet(b)
P("Resultados cuantitativos destacados:", bold=True)
for b in ["Índices en consultas filtradas: hasta 97,7 % de mejora (Seq Scan → Index Scan).",
          "Particionamiento (partition pruning): 79,6 % (15,07 → 3,07 ms).",
          "Reescritura de ventas por categoría: 50,8 % (15,68 → 7,72 ms).",
          "MongoDB, índice compuesto ESR: docsExamined 1.000 → 198 (−80,2 %), 18,4 → 1,8 ms.",
          "MongoDB, pipeline analítico optimizado: 84,9 % (284 → 43 ms)."]:
    bullet(b)

# ===================== b =====================
H("b. Implementación PostgreSQL")
H("b.1 Scripts DDL y volumen real", 2)
P("Esquema relacional con la estructura real de Olist (8 tablas) con PK, FK y constraints. "
  "DDL completo en postgresql/schema/. La tabla orders se migró a un esquema particionado "
  "(orders_part) por RANGE de order_purchase_timestamp.")
table(["Tabla","Registros","Tamaño"],[
 ["orders","100.000","13 MB"],["order_items","15.000","1.760 kB"],["order_payments","5.000","528 kB"],
 ["customers","1.000","144 kB"],["products","500","104 kB"],["categories / sellers / geolocations","20 / 100 / 100","—"]])
P("Estado de order_status: delivered 69.516 · shipped 29.424 · processing 1.060 (1,06 %).")

H("b.2 Estrategia de indexación justificada", 2)
table(["Tipo de consulta / filtro","Índice","Justificación técnica"],[
 ["Llaves foráneas y JOINs recurrentes (customer_id, order_id, product_id, seller_id, category_id)","B-Tree","Columnas usadas frecuentemente en Hash Joins; B-Tree es de bajo costo y alta eficiencia."],
 ["Estado específico u operaciones anómalas (p. ej. 'processing')","Parcial","Optimiza búsquedas críticas de un subconjunto pequeño sin ocupar el espacio del resto de la tabla."],
 ["Filtros geográficos o regionales","B-Tree","Agiliza las agrupaciones de métricas de ventas por ubicación."],
 ["Rangos de fecha de gran volumen","BRIN","Índice diminuto orientado a la escalabilidad del histórico de órdenes."],
 ["Atributos no estructurados (JSONB)","No aplica","El esquema relacional no usa campos JSONB; no se requieren índices GIN en esta entrega."],
 ["Búsquedas textuales complejas","No aplica","Las consultas se basan en relaciones y filtros exactos; no hay full-text."]])

H("b.3 Índices creados", 2)
table(["Índice","Tipo","Tamaño","Patrón que optimiza"],[
 ["idx_orders_customer_id","B-Tree","704 kB","JOIN/filtro por customer_id"],
 ["idx_order_items_product_id","B-Tree","128 kB","JOIN/filtro por product_id"],
 ["idx_order_items_seller_id","B-Tree","120 kB","JOIN/filtro por seller_id"],
 ["idx_order_items_seller_incl_price","B-Tree + INCLUDE","480 kB","Agregación de price (Index Only Scan)"],
 ["idx_orders_processing","Parcial","40 kB","Pedidos 'processing' (1 % de la tabla)"],
 ["idx_orders_purchase_brin","BRIN","24 kB","Rangos de fecha (escalabilidad)"]])

H("b.4 Particionamiento aplicado", 2)
for b in ["Tabla: orders (100.000 filas) · Columna: order_purchase_timestamp",
          "Tipo: RANGE · Granularidad: mensual (14 particiones + DEFAULT) · índices locales por partición",
          "Cada partición pesa ~870 kB con ~8.000 filas, manteniendo bloques predecibles para los reportes.",
          "Mantenimiento: retención 24 meses, archivado vía DETACH PARTITION hacia cold storage, creación automática de la partición del mes siguiente con pg_cron."]:
    bullet(b)
P("Validación — partition pruning:", bold=True)
table(["Escenario","Exec (ms)","Relaciones escaneadas"],[
 ["Sin particionamiento","15,07","orders (tabla completa)"],
 ["Con particionamiento","3,07","orders_part_202512 (1 partición)"]])
P("Mejora: 79,6 % — el planificador descarta 13 particiones y solo lee una.")
figura(os.path.join(EXPL,"particion_pruning_explain.png"),
       "Figura. EXPLAIN (ANALYZE, BUFFERS) de la consulta por rango sobre orders_part (partition pruning).")

H("b.5 Queries críticas optimizadas", 2)
P("Se priorizaron 5 consultas por frecuencia e impacto en el negocio (Q1–Q5). Sobre ellas se "
  "aplicaron tres técnicas de reescritura, cada una justificada por su plan de ejecución:")
table(["Técnica","Consulta","Antes","Después","Mejora"],[
 ["1. Pre-agregación en CTE","Q1 Ventas por categoría","15,68 ms","7,72 ms","50,8 %"],
 ["2. Corrección de fan-out con CTEs","Q5 Reporte de órdenes","60,69 ms","48,48 ms","20,1 %"],
 ["3. Eliminación de JOIN redundante","Q3 Clientes top","29,97 ms","28,25 ms","5,7 %"]])
P("Técnica 1 — Pre-agregación en CTE (Q1): agregar order_items por producto antes de unir con "
  "products y categories reduce las filas que entran al JOIN (15.000 → ~500).")
figura(os.path.join(EXPL,"q1_cte_antes.jpeg"),  "Figura 4a. EXPLAIN de Q1 original (Seq Scan).")
figura(os.path.join(EXPL,"q1_cte_despues.jpeg"),"Figura 4b. EXPLAIN de Q1 reescrita con CTE (8,6 ms).")
P("Técnica 2 — Corrección de fan-out con CTEs (Q5): Q5 unía order_items y order_payments al mismo "
  "order_id, multiplicando filas (ítems × pagos) e inflando los SUM. Pre-agregar cada lado en su "
  "propia CTE corrige el doble conteo: mejora de rendimiento y de correctitud a la vez.")
figura(os.path.join(EXPL,"q5_fanout_antes.jpeg"),  "Figura 5a. EXPLAIN de Q5 original (fan-out).")
figura(os.path.join(EXPL,"q5_fanout_despues.png"), "Figura 5b. EXPLAIN de Q5 reescrita con CTEs.")
P("Técnica 3 — Eliminación de JOIN redundante (Q3): Q3 unía customers solo para leer customer_id, "
  "columna que ya existe en orders; eliminar ese JOIN quita una tabla del plan sin cambiar el resultado.")
figura(os.path.join(EXPL,"q3_join_redundante.jpeg"),"Figura 6. EXPLAIN de Q3 original (3 tablas) vs. reescrita (2 tablas).")

# ===================== c =====================
H("c. Implementación MongoDB")
H("c.1 Colecciones, esquemas y validación", 2)
P("El módulo analítico se implementó en MongoDB Atlas con dos colecciones. Cada colección refleja "
  "cómo la aplicación consulta los datos: los productos embeben sus especificaciones variables y "
  "métricas precalculadas; las reseñas se mantienen separadas porque su ciclo de escritura es "
  "independiente del catálogo.")
table(["Colección","Docs","Patrón","Decisión de modelado"],[
 ["products","1.000","Attribute/Embedding + Computed","Specs embebidas (varían por categoría) y métricas precalculadas (total_units_sold, average_rating) para evitar aggregations en cada lectura"],
 ["order_reviews","500","Referenced","Escritura alta e independiente del catálogo; embeber generaría documentos de crecimiento ilimitado"]])
P("Documento de ejemplo (products) con Attribute Pattern + Computed Pattern:", bold=True)
code(
"""{
  "_id": "PROD-0001",
  "name": "Producto Computers 1",
  "category": { "id": 1, "name": "Computers" },
  "price": { "amount": 1245.50, "currency": "USD" },
  "seller_id": "SELL-023",
  "status": "ACTIVE",
  "specifications": { "processor": "Intel i5", "ram": "16GB", "storage": "512GB SSD" },
  "computed_metrics": {
    "total_units_sold": 312, "average_rating": 4.3,
    "total_reviews": 87, "last_updated": "2026-06-01T00:00:00Z"
  }
}""")
P("Ambas colecciones se crean con validación JSON Schema ($jsonSchema, validationAction:\"error\"); "
  "el código completo de los validadores está en mongodb/schema/01_collections_schema.js.")

H("c.2 Índices (regla ESR) y validación", 2)
P("Se implementaron cuatro índices siguiendo la regla ESR (Equality, Sort, Range):")
table(["Índice","Tipo","Colección","Optimiza"],[
 ["idx_products_category_price_status","Compuesto ESR","products","Catálogo por categoría ordenado por precio"],
 ["idx_products_active","Parcial (status=ACTIVE)","products","Reduce el índice ~80 %; excluye ~200 productos inactivos"],
 ["idx_products_text_name","Text","products","Búsqueda full-text por nombre"],
 ["idx_reviews_product_score","Compuesto","order_reviews","Reseñas de un producto por calificación"]])
P("Validación con .explain(\"executionStats\"):", bold=True)
table(["Índice","docsExamined antes","docsExamined después","Tiempo antes","Tiempo después"],[
 ["idx_products_category_price_status","1.000","198","18,4 ms","1,8 ms"],
 ["idx_products_active","1.000","802","12,7 ms","1,2 ms"],
 ["idx_reviews_product_score","500","43","8,9 ms","0,7 ms"]])
P("El compuesto ESR logra la mayor mejora: docsExamined 1.000 → 198 (−80,2 %), porque el "
  "planificador usa categoría + precio para localizar el subconjunto relevante sin recorrer la colección.")

H("c.3 Aggregation pipeline (7 stages)", 2)
P("Pipeline de análisis de rendimiento del catálogo: $match (índice parcial) → $lookup (reviews) → "
  "$unwind → $group (métricas por categoría) → $addFields (etiqueta de rendimiento) → $sort → "
  "$project, con allowDiskUse: true. Código en mongodb/schema/03_aggregation_pipelines.js.")
table(["Escenario","executionTimeMillis","docsExamined","Observación"],[
 ["Sin $match inicial ni índices","284 ms","1.000","Recorre toda la colección antes del lookup"],
 ["Con índice parcial pero $match al final","187 ms","1.000","El índice existe pero el planificador no lo aprovecha"],
 ["Con $match al inicio e índices activos","43 ms","198","−84,9 % respecto al escenario sin optimizar"]])
P("El $match en Stage 1 aprovecha el índice parcial desde el primer paso (1.000 → 802 docs antes del "
  "$lookup); el $project al final no bloquea el uso de índices en stages anteriores.")

H("c.4 Diseño teórico de sharding y replica sets", 2)
P("Sharding — colección candidata: analytics_events/orders. Shard key recomendada: "
  "{ customer_id: 1, order_purchase_timestamp: 1 } (range sobre clave compuesta).")
for b in ["customer_id aporta alta cardinalidad y dispersa las escrituras (evita el hot shard de una clave temporal monotónica).",
          "order_purchase_timestamp conserva la localidad de rango → analíticas por cliente+periodo targeted, no scatter-gather.",
          "Hashed vs range: se descarta hashed porque convertiría todo reporte temporal en scatter-gather.",
          "Distribución simulada (3 shards por rangos de customer_id): ~33 % / 33 % / 34 %."]:
    bullet(b)
P("Replica sets — topología PSS (Primary + 2 Secondary) multi-AZ por shard. 3 miembros con datos "
  "garantizan quórum 2/3 para w:majority y tolerancia a la caída de 1 nodo (se descarta Arbiter).")
table(["Operación","Write Concern","Read Concern","Read Preference"],[
 ["Transaccional (orden/pago/reseña)","majority, j:true","majority","primary"],
 ["Lectura analítica (reportes)","—","local","secondaryPreferred"],
 ["Baja criticidad (logs/telemetría)","w:1","local","primaryPreferred"],
 ["Read-your-writes","majority","majority","primary / sesión causal"]])
figura(os.path.join(EJ3,"Diagrama_ReplicaSet_3nodos_multiAZ.jpg"),
       "Figura. Replica Set MongoDB de 3 nodos en multi-AZ (Ecommify).")

# ===================== d =====================
H("d. Evidencias cuantitativas de mejoras de rendimiento")
H("d.1 PostgreSQL — líneas base (solo índices de PK)", 2)
table(["Consulta","Exec (ms)","Tablas con Seq Scan"],[
 ["Q1 Ventas por categoría","87,89","order_items, products, categories"],
 ["Q2 Vendedores por ingresos","10,87","order_items, sellers"],
 ["Q3 Clientes mayores compras","52,88","orders, order_payments, customers"],
 ["Q4 Ventas por estado","62,64","order_items, customers, geolocations, products, categories"],
 ["Q5 Reporte de órdenes","55,79","order_items, orders, order_payments, customers"]])

H("d.2 Impacto de índices (consultas filtradas)", 2)
table(["Consulta","Antes","Después","Plan","Mejora"],[
 ["F1 customer_id","10,15 ms","0,23 ms","Seq → Bitmap Heap Scan","97,7 %"],
 ["F2 seller_id","1,64 ms","0,24 ms","Seq → Bitmap Heap Scan","85,6 %"],
 ["F3 status='processing'","9,91 ms","0,28 ms","Seq → Index Only Scan","97,1 %"],
 ["F4 rango 1 mes (BRIN)","11,12 ms","10,65 ms","Seq → Seq","4,3 %"]])
figura(os.path.join(EVID,"01_indices_consultas_filtradas.png"),
       "Figura. PostgreSQL — impacto de índices especializados en consultas filtradas.")

H("d.3 Optimización de queries críticas", 2)
figura(os.path.join(EVID,"02_optimizacion_queries_criticas.png"),
       "Figura. PostgreSQL — optimización de queries críticas (reescritura SQL).")

H("d.4 Particionamiento (partition pruning)", 2)
figura(os.path.join(EVID,"03_particionamiento_pruning.png"),
       "Figura. PostgreSQL — particionamiento por RANGE (partition pruning, 79,6 %).")

H("d.5 MongoDB — índices y aggregation pipeline", 2)
figura(os.path.join(EVID,"04_mongodb_indices.png"),
       "Figura. MongoDB — impacto de índices (executionTimeMillis, .explain).")
figura(os.path.join(EVID,"05_mongodb_pipeline.png"),
       "Figura. MongoDB — aggregation pipeline analítico optimizado (84,9 %).")

H("d.6 Interpretación", 2)
for b in ["Las mayores mejoras provienen de decisiones de diseño (índice parcial, pre-agregación, partition pruning, regla ESR), no de fuerza bruta.",
          "El BRIN mejora poco a 100k filas cacheadas pero se justifica por su costo ínfimo (24 kB) y escalabilidad.",
          "En MongoDB, ubicar $match al inicio y $project al final permite aprovechar el índice parcial desde el primer stage."]:
    bullet(b)

# ===================== e =====================
H("e. Sincronización entre sistemas y estrategia de consistencia")
P("PostgreSQL es la fuente de verdad transaccional (CP); MongoDB es un read model analítico (AP) "
  "derivado y enriquecido. Flujo unidireccional para datos maestros mediante un job ETL programado "
  "que hace upsert por _id en MongoDB (idempotente). Las reseñas y eventos nacen en MongoDB.")
P("Consistencia eventual — doble ventana de inconsistencia:", bold=True)
P("Lag percibido = Ventana 1 (latencia del job ETL: segundos–minutos) + Ventana 2 (replication lag: ~ms–2 s).")
table(["Caso de uso","Estrategia de mitigación"],[
 ["Operación crítica (ver último estado)","Leer de PostgreSQL o readPreference:primary + readConcern:majority"],
 ["Read-your-writes","Sesión con causal consistency"],
 ["Evitar leer datos que podrían perderse en rollback","readConcern:majority + writeConcern:majority"],
 ["Lectura tolerante (catálogo, reportes)","secondaryPreferred/nearest, readConcern:local"]])
figura(os.path.join(EJ3,"Diagrama_Consistencia_Eventual_ventana_lag.jpg"),
       "Figura. Consistencia eventual en Ecommify: doble ventana de inconsistencia.")

# ===================== f =====================
H("f. Lecciones aprendidas")
P("Obstáculos y soluciones:", bold=True)
for b in ["Datos reales sin atributos semiestructurados → se justificó técnicamente no forzar JSONB/GIN en PostgreSQL y concentrar la flexibilidad documental en MongoDB (patrón Attribute/Embedding).",
          "Cifras desactualizadas (5.000 vs. 100.000 órdenes) → reejecución de mediciones; particionamiento justificado por volumen real.",
          "Seq Scan óptimo en agregaciones de tabla completa → impacto de índices demostrado con consultas filtradas (85–98 %).",
          "Riesgo de fan-out en Q5 (ítems × pagos) → corregido con CTEs por lado, mejorando rendimiento y correctitud.",
          "Datos derivados para MongoDB → derivación documentada (no presentar datos inventados como reales)."]:
    bullet(b)
P("Limitaciones del free tier y workarounds:", bold=True)
table(["Limitación","Workaround"],[
 ["Atlas M0 sin sharding real","Sharding documentado de forma teórica (shard key, hashed vs range, simulación)"],
 ["M0 sin control fino del replica set","Topología PSS multi-AZ como diseño objetivo"],
 ["Supabase: DDL pesado / pg_cron limitado","Session Pooler (5432) para DDL; mantenimiento de particiones documentado"],
 ["Pipelines que exceden memoria","allowDiskUse: true"]])
P("Conclusión: la arquitectura políglota demostró su valor: PostgreSQL para integridad transaccional "
  "y MongoDB para analítica/catálogo flexible. Trabajar con datos reales obligó a decisiones honestas "
  "de modelado, más valiosas que un esquema ideal con datos ficticios.")

# ===================== 2 y 3 =====================
H("2. Repositorio GitHub")
P("Estructura organizada (postgresql/, mongodb/, notebooks/, docs/, evidencias/), README con "
  "instrucciones de setup para Supabase y MongoDB Atlas, scripts idempotentes y reproducibles. "
  "El módulo MongoDB incluye los validadores JSON Schema, los índices ESR y el aggregation pipeline "
  "en mongodb/schema/.")

H("3. Video de demostración")
P("El video (5–10 min) recorre la conexión a Supabase y Atlas, la navegación por tablas/colecciones, "
  "la ejecución de queries optimizadas mostrando las mejoras y la explicación de las decisiones "
  "técnicas clave (índice parcial 'processing', partition pruning, regla ESR en MongoDB).")
P("Enlace al video:", bold=True)
pv = doc.add_paragraph(); rv = pv.add_run(VIDEO_URL); rv.font.size=Pt(10); rv.font.color.rgb=AZUL; rv.underline=True

doc.save(OUT)
print("DOCX generado:", OUT)
